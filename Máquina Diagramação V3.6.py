import docx
import datetime
import os

# Caminho do arquivo
caminho_arquivo = r'caminho_entrada'
# Extrair o nome do arquivo do caminho
nome_arquivo_original = os.path.basename(caminho_arquivo)

print(nome_arquivo_original)

# Abrir o arquivo docx
doc = docx.Document(caminho_arquivo)

# Criar variável para armazenar o índice do parágrafo
indice = 0
# Criar variável para armazenar os índices dos parágrafos a serem removidos
remover = []

# Percorrer todas as linhas do arquivo
for paragraph in doc.paragraphs:
    # Verificar se a linha está em CAPS LOCK
    if paragraph.text.isupper():
        # Adicionar uma linha em branco após o parágrafo em CAPS LOCK
        #--doc.paragraphs[indice].add_run("\n")
        # Aplicar negrito
        for run in paragraph.runs:
            run.font.bold = True
        # Aplicar avanço de parágrafo caso ele já não exista
        if paragraph.paragraph_format.left_indent is None:
            paragraph.paragraph_format.left_indent = docx.shared.Inches(0.5)
    # Alterar a fonte do texto inteiro para "Quattrocento Sans"
    for run in paragraph.runs:
        run.font.name = 'Quattrocento Sans'
    # Incrementar o índice em 1
    indice += 1

# Salvar o arquivo com o nome original acrescido da data e hora de criação
data_e_hora = datetime.datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
caminho_saida = 'caminho_saida'
nome_arquivo_sem_extensao = os.path.splitext(nome_arquivo_original)[0]
nome_arquivo_novo = f'[MF 3.6]{nome_arquivo_sem_extensao}_{data_e_hora}.docx'
caminho_completo = os.path.join(caminho_saida, nome_arquivo_novo)
doc.save(caminho_completo)

# Imprimir uma mensagem de sucesso
print(f"Arquivo modificado salvo como {nome_arquivo_novo}")
